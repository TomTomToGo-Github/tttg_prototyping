import docx
from docx.shared import Mm
from docx import Document
# Assuming 'doc' is a Document object
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
doc = Document('test.docx')

import xml.etree.ElementTree as ET


import xml.etree.ElementTree as ET
7
        # Find and extract specified anchor attributes
        # Extract the desired XML segment
        # print(xml_str)
        # start_string = '<w:r xmlns'
        # end_string = '</w:r>'
        start_string = '<w:drawing>'
        # start_string = '<w:drawing>'
        # end_string = '</w:drawing>'
        end_string = '<pic:nvPicPr>'
        start_index = xml_str.find(start_string)
        end_index = xml_str.find(end_string) + len(end_string)
        if start_index != -1 and end_index != -1:
            extracted_xml = xml_str[start_index:end_index]
            formatting['image_xml'] = extracted_xml
            # print(extracted_xml)
        else:
            print("XML segment not found.")
        
        # # Find wp:simplePos and extract x and y attributes
        # for simple_pos in xml_tree.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}simplePos'):
        #     x = simple_pos.attrib.get('x')
        #     y = simple_pos.attrib.get('y')
        #     print(f"SimplePos: x={x}, y={y}")


        # Get the XML element for the image
        # print(dir(picture))
        # print(dir(picture._inline))
        # print(new_image)
        if 'image_xml' in new_image:
            # graphic_frame = picture._inline.graphic.graphicData.pic.blipFill.blip.getparent().getparent().getparent()
            # Replace the inline picture with the anchored picture
            # print(graphic_frame)
            print('\n\n\n')
            print('\n\n\n')
            # print('\n\n\n', new_image['image_xml'])
            # print(parse_xml(new_image['image_xml']))
            # print(picture._inline.getparent().getparent().xml)
            current_xml = picture._inline.getparent().getparent().xml
            start_string = '<w:drawing>'
            end_string = '<pic:nvPicPr>'
            # start_string = '<w:drawing>'
            # end_string = '</w:drawing>'
            start_index = current_xml.find(start_string)
            end_index = current_xml.find(end_string) + len(end_string)
            if start_index != -1 and end_index != -1:
                new_xml = new_image['image_xml']
                print("\n\n Current:\n" + current_xml[:start_index] )
                print("\n\n Selection: \n", current_xml[start_index:end_index])
                print("\n\n Original with anchor:\n" + new_xml)
                current_xml = current_xml.replace(current_xml[start_index:end_index], new_image['image_xml'])
                current_xml = current_xml.replace('</wp:inline>', '</wp:anchor>')
                print("\n\nOriginal xml:\n" + picture._inline.getparent().getparent().xml)
                print("\n\nReplaced xml:\n" + current_xml)
                picture._inline.getparent().getparent().replace(parse_xml(picture._inline.getparent().getparent().xml), parse_xml(current_xml))
                print("\n\Replaced2 xml:\n" + picture._inline.getparent().getparent().xml)

                # picture._inline.getparent().getparent().replace(parse_xml(current_xml[start_index:end_index]), parse_xml(new_image['image_xml']))
            else:
                print("XML segment not found.")

            print('\n\n\n')
            print('\n\n\n')
            print('\n\n\n')

        #     picture._inline.getparent().replace(picture._inline, parse_xml(new_image['image_xml']))     
            # anchor = picture._inline.getparent().getparent()            
            
            
            # for key, value in new_image['anchor'].items():
            #     anchor.set(key, value or '1')
        # anchor.set('simplePos', new_imag
        # e.get('simplePos','1'))
        # anchor.set('relativeHeight', new_image.get('relativeHeight','1'))
        # anchor.set('behindDoc', new_image.get('behindDoc','0'))
        # anchor.set('locked', new_image.get('locked','0'))
        # anchor.set('layoutInCell', new_image.get('layoutInCell','1'))
        # anchor.set('allowOverlap', new_image.get('allowOverlap','1'))






# Your XML snippet
xml_snippet = '''
<w:drawing>
    <wp:anchor behindDoc="0" distT="0" distB="0" distL="114935" distR="114935" simplePos="0" locked="0" layoutInCell="0" allowOverlap="1" relativeHeight="3">
        <wp:simplePos x="0" y="0"/>
        <wp:positionH relativeFrom="column">wp:posOffset591820</wp:positionH>
        <wp:positionV relativeFrom="paragraph">wp:posOffset76200</wp:positionV>
        <wp:extent cx="3827145" cy="695960"/>
        <wp:effectExtent l="0" t="0" r="0" b="0"/>
        <wp:wrapSquare wrapText="largest"/>
        <wp:docPr id="2" name="Image1" descr=""/>
        <wp:cNvGraphicFramePr>
            <a:graphicFrameLocks xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" noChangeAspect="1"/>
        </wp:cNvGraphicFramePr>
        <a:graphic xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">
            <a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">
                <pic:pic xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture">
                    <pic:nvPicPr>
                        <pic:cNvPr id="2" name="Image1" descr=""/>
                        <pic:cNvPicPr>
                            <a:picLocks noChangeAspect="1" noChangeArrowheads="1"/>
                        </pic:cNvPicPr>
                    </pic:nvPicPr>
                </pic:pic>
            </a:graphicData>
        </a:graphic>
    </wp:anchor>
</w:drawing>
'''

# Parse the XML snippet
root = ET.fromstring(xml_snippet)

# Find wp:simplePos and extract x and y attributes
for simple_pos in root.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}simplePos'):
    x = simple_pos.attrib.get('x')
    y = simple_pos.attrib.get('y')
    print(f"SimplePos: x={x}, y={y}")
7

'relativeHeight', 'behindDoc', 'locked', 'layoutInCell', 'allowOverlap'








ns = {
    'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
    'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
}
element_object_mapping = {'paragraphs': {}, 'tables': {}, 'images': {}, 'images_floating': {}, 'others': {}}
paragraph_iter = iter(doc.paragraphs)
table_iter = iter(doc.tables)
image_iter = iter_and_save_images_from_docx(doc)
for element in doc.element.body:
    if element.tag.endswith('p'):
        paragraph = next(paragraph_iter)
        element_object_mapping['paragraphs'][element] = paragraph
        for run in paragraph.runs:
            if 'graphicData' in run._r.xml:
                last_xml = run._r.xml
                print(last_xml)
                # Extract image size information
                xml_tree = ET.fromstring(run._r.xml)
                for elem in xml_tree.iter(f'{{{ns["a"]}}}ext'):
                    cx = int(elem.attrib.get('cx'))
                    cy = int(elem.attrib.get('cy'))
                    print(f"Image size: Width: {cx / 914400}, Height: {cy / 914400}")
                print("Link elements: Found image in paragraph run")
        

























# Load the Word document (a .docx file is a zip file)
docx_file = 'test.docx'

# Extract the XML content of the document
with zipfile.ZipFile(docx_file, 'r') as z:
    xml_content = z.read('word/document.xml')

# Parse the XML content
xml_tree = ET.fromstring(xml_content)

# Namespace mapping for XML elements
ns = {
    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
    'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture'
}

# Find image elements and get their size
for elem in xml_tree.iter(f'{{{ns["w"]}}}r'):
    for pic_elem in elem.iter(f'{{{ns["pic"]}}}pic'):
        for shape_elem in pic_elem.iter(f'{{{ns["pic"]}}}spPr'):
            for cx_elem in shape_elem.iter(f'{{{ns["pic"]}}}cx'):
                cx = int(cx_elem.attrib.get('val'))
                cy = int(shape_elem.find(f'{{{ns["pic"]}}}cy').attrib.get('val'))
                print(f"Image size: Width: {cx}, Height: {cy}")


























body = doc._element.body
# Find the floating image
for element in body.iter():
    if 'drawing' in element.tag:
        print("Link elements: Found image as floating")
        # Find the horizontal position element
        for posH in element.iter():
            if 'posH' in posH.tag:
                # This is the horizontal position element
                # Change the position
                posH.set('dist', '720')  # 720 is in EMUs (1 inch = 914400 EMUs)
# Save the document
doc.save('output.docx')


for i, image in enumerate(doc.inline_shapes):
    original_width, original_height = image.width, image.height
    
def get_temp_img_name(num):
    return f'__temp__saved_image_{num}.png'

def iter_and_save_images_from_docx(doc):
    doc_rels = doc.part.rels
    ctr_img_read = 0
    for rel in doc_rels.values():
        if "image" in rel.reltype:
            ctr_img_read += 1
            image_data = rel._target._blob
            image_path = get_temp_img_name(ctr_img_read)
            with open(image_path, 'wb') as img_file:
                img_file.write(image_data)
            yield image_path, ctr_img_read





    def link_elements_to_objects(doc):
        element_object_mapping = {'paragraphs': {}, 'tables': {}, 'images': {}, 'images_floating': {}, 'others': {}}
        paragraph_iter = iter(doc.paragraphs)
        table_iter = iter(doc.tables)
        image_iter = iter_and_save_images_from_docx(doc)
        for element in doc.element.body:
            if element.tag.endswith('p'):
                try:
                    paragraph = next(paragraph_iter)
                    element_object_mapping['paragraphs'][element] = paragraph
                    for run in paragraph.runs:
                        if 'graphicData' in run._r.xml:
                            print("Link elements: Found image in paragraph run")
                            image_local_path, ctr_img_read = next(image_iter)
                            element_object_mapping['images'][ctr_img_read] = image_local_path
                except StopIteration:
                    break
            elif element.tag.endswith('tbl'):
                try:
                    table = next(table_iter)
                    element_object_mapping['tables'][element] = table
                except StopIteration:
                    break
            elif 'graphicData' in element.xml:
                print("Link elements: Found image as stand - allone")
                try:
                    image_local_path, ctr_img_read = next(image_iter)
                    element_object_mapping['images'][ctr_img_read] = image_local_path
                except StopIteration:
                    break
            elif element.tag.endswith('secPr'):
                print("Link elements: section properties found")
                break
            else:
                print("Link elements: Something else found")
                element_object_mapping['others'][element] = "??"
                break

        for image_local_path, ctr_img_read in image_iter:
            print("Link elements: Found more images - probably floating")
            # image_local_path, ctr_img_read = next(image_iter)
            element_object_mapping['images_floating'][ctr_img_read] = image_local_path





# def main_test_document_generation():
#     document = Document()
#     document.add_heading('Document Title', 0)

#     p = document.add_paragraph('A plain paragraph having some ')
#     p.add_run('bold').bold = True
#     p.add_run(' and some ')
#     p.add_run('italic.').italic = True

#     document.add_heading('Heading, level 1', level=1)
#     document.add_paragraph('Intense quote', style='Intense Quote')
#     document.add_paragraph('first item in unordered list', style='List Bullet')
#     document.add_paragraph('first item in ordered list', style='List Number')
#     document.add_picture('test.png', width=Inches(1.25))
#     records = (
#         (3, '101', 'Spam'),
#         (7, '422', 'Eggs'),
#         (4, '631', 'Spam, spam, eggs, and spam')
#     )

#     table = document.add_table(rows=1, cols=3)
#     hdr_cells = table.rows[0].cells
#     hdr_cells[0].text = 'Qty'
#     hdr_cells[1].text = 'Id'
#     hdr_cells[2].text = 'Desc'
#     for qty, id, desc in records:
#         row_cells = table.add_row().cellxs
#         row_cells[0].text = str(qty)
#         row_cells[1].text = id
#         row_cells[2].text = desc
#     document.add_page_break()

#     document.save('demo.docx')
    
    def link_elements_to_objects(doc):
        element_object_mapping = {'paragraphs': {}, 'tables': {}, 'images': {}}
        paragraph_iter = iter(doc.paragraphs)
        table_iter = iter(doc.tables)
        image_iter = iter(doc.inline_shapes)
        for element in doc.element.body:
            if element.tag.endswith('p'):
                try:
                    paragraph = next(paragraph_iter)
                    element_object_mapping['paragraphs'][element] = paragraph
                except StopIteration:
                    break
            elif element.tag.endswith('tbl'):
                try:
                    table = next(table_iter)
                    element_object_mapping['tables'][element] = table
                except StopIteration:
                    break
            elif 'graphicData' in element.xml:
                try:
                    print('found image')
                    image = next(image_iter)
                    element_object_mapping['images'][element] = image
                except StopIteration:
                    break
        return element_object_mapping

    def scan_docx(filename):
        doc = docx.Document(filename)
        paragraph_count = 0
        for element in doc.element.body:
            element.tag
            if element.tag.endswith('p'):
                paragraph_count += 1
                if not element.style and paragraph_count > 2:
                    print(element.style)
                    print(element.runs)
                    dir(element)
                
                print(element.style.name)
                print(element.text)
            elif element.tag.endswith('tbl'):
                print(f'Table found after {paragraph_count} paragraphs')
            elif element.tag.endswith('drawing'):
                print(f'Image found after {paragraph_count} paragraphs')

    for para in doc.paragraphs:
        new_element = None

    for shape in doc.inline_shapes:
        print(shape)

    doc = docx.Document(filename)
    paragraph_count = 0
    image_count = 0
    for index, paragraph in enumerate(doc.paragraphs):
        if paragraph.text:
            paragraph_count += 1
        for run in paragraph.runs:
            :
                image_count += 1
                print(f'Image found after {paragraph_count} paragraphs')
    print(f'Total images: {image_count}')

def scan_docx(filename):
    doc = docx.Document(filename)
    for tab in doc.tables:
        print(tab)
        
    for shape in doc.inline_shapes:
        print(shape)
    
    
    
    for st in doc.styles:
        print(st.name)
    font = doc.styles['Normal'].font
    font.name
    font.size
    
    section = doc.sections[0]
    print(section.page_height)
    print(section.page_width)
    print(section.left_margin)
    print(section.right_margin)
    print(section.top_margin)
    print(section.bottom_margin)
    print(section.header_distance)
    print(section.footer_distance)
    # section.page_height = Mm(297)
    # section.page_width = Mm(210)
    # section.left_margin = Mm(25.4)
    # section.right_margin = Mm(25.4)
    # section.top_margin = Mm(25.4)
    # section.bottom_margin = Mm(25.4)
    # section.header_distance = Mm(12.7)
    # section.footer_distance = Mm(12.7)
    full_doc = list()
    for para in doc.paragraphs:
        for para in doc.paragraphs:
            new_element = None
            if para.style.name == 'Title':
                new_element = {
                    "type": "header",
                    "level": 0,
                    "content": para.text,
                    # "body": []
                }
            elif 'Heading' in para.style.name:
                new_element = {
                    "type": "header",
                    "level": int(para.style.name[-1]),
                    "content": para.text,
                    # "body": []
                }
            elif para.style.name in ['Normal', 'Intense Quote', 'List Bullet', 'List Number']:
                if len(para.runs) == 1 and para.style.name != 'Normal':
                    content = para.text
                elif len(para.runs) >= 1 and para.style.name == 'Normal':
                    content = []
                    for run in para.runs:
                        new_run = {
                            "type": "paragraph",
                            "style": 'bold' if run.bold else 'italic' if run.italic else 'Normal',
                            "content": run.text
                        }
                        content.append(new_run)
                else:
                    print("Unhandled paragraph case!!!")
                    content = 'UNHANDLED CASE'
                new_element = {
                    "type": "paragraph",
                    "style": para.style.name,
                    "content": content,
                }
        full_doc.append(new_element)

    return full_doc
    # ['Title', 'Normal', 'Normal', 'Heading 1', 'Intense Quote', 'List Bullet', 'List Number', 'Normal']
    # return '\n'.join(fullText)

filename = "test.docx"
scan_docx(filename)

from docx import Document
import json

document = Document('test.docx')
lines = [para.text for para in document.paragraphs]
lines = [line.partition('.') for line in lines]
lines = [(int(row_num), row_text) for row_num, _, row_text in lines]
lines = [(n, [txt.partition(':') for txt in row_text.split(',')]) for n, row_text in lines]
lines = {n: {key.strip(): val.strip() for key, _, val in row} for n, row in lines}
json_result = json.dumps(lines)


import docx2txt, json, collections
# step 1 get docx text
text = docx2txt.process("test.docx")
# convert to list
li = [x for x in text.split('\n')]
# remove ''s i.e Nones
li = list(filter(None, li))
print(li)
# json list
json_li = []
# convert and store all values
for x in li:
    x = x[2:] # remove 1. 2. 3. ...
    y = x.split(',')
    print(y)
    d = collections.defaultdict()
    for m in y:
        z = m.split(':')
        z1 = [x.strip() for x in z]
        d[z1[0]] = z1[1]
    json_li.append(d)
# JSON conversion
print(json.dumps(json_li, indent=4))